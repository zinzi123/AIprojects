#comb of imp2.py + new1.py 

import streamlit as st
import openai
from dotenv import load_dotenv
import os
from docx import Document
from io import BytesIO

# Load environment variables from .env file
load_dotenv()

# Set your OpenAI API key from environment variables
openai.api_key = os.getenv("OPENAI_API_KEY")

# Title of the app
st.title("Automated SOW Generator")

# Collecting minimal user inputs
st.header("Enter Basic Project Details")
project_title = st.text_input("Project Title")
client_name = st.text_input("Client Name")
project_description = st.text_area("Project Description")
start_date = st.date_input("Start Date")
end_date = st.date_input("End Date")
budget = st.text_input("Budget")

# Select the type of SOW
sow_type = st.selectbox(
    "Select the type of SOW",
    ["Fixed Bid", "T&M (Time & Materials)", "Managed Capacity", "Managed Services / Support"]
)

# Role and rates input for T&M and Managed Capacity
roles_involved = []
if sow_type in ["T&M (Time & Materials)", "Managed Capacity"]:
    st.subheader("Roles and Rates")
    num_roles = st.number_input("Number of Roles", min_value=1, step=1, value=1)
    for i in range(num_roles):
        role = st.text_input(f"Role {i + 1}")
        rate = st.text_input(f"Rate for {role} (per hour/month)")
        if role and rate:
            roles_involved.append((role, rate))

# Define headings for each SOW type
sow_headings = {
    "Fixed Bid": [
        "Legal Entity", "Description of Project", "Scope", "Services Included", "Deliverables", 
        "Start and End Date", "Milestone", "Payment / Invoice Schedule", "Upfront 35% Invoicing", 
        "IP or Licenses Secured/Procured by G10X", "Passthrough Expenses", "Governance Model", 
        "Out of Scope", "Assumptions", "Dependencies", "Warranty", "Early Exit / Termination Clause", 
        "Change Management", "Signatory (based on location)"
    ],
    "T&M (Time & Materials)": [
        "Legal Entity", "Description of Project", "Scope", "Services Included", "Deliverables", 
        "Start and End Date", "Team Structure (Location & Headcount)", "Rate Card", 
        "Payment / Invoice Schedule", "Upfront 35% Invoicing", "IP or Licenses Secured/Procured by G10X", 
        "Passthrough Expenses", "Governance Model", "Out of Scope", "Assumptions", "Dependencies", 
        "Change Management", "Warranty", "Early Exit / Termination Clause", "Signatory (based on location)"
    ],
    "Managed Capacity": [
        "Legal Entity", "Description of Project", "Scope", "Services Included", "Deliverables", 
        "Start and End Date", "Team Structure (Location & Headcount)", "Metrics, KPIs", "Rate Card", 
        "Capacity Fee", "Payment / Invoice Schedule", "Upfront 35% Invoicing", 
        "IP or Licenses Secured/Procured by G10X", "Passthrough Expenses", "Governance Model", 
        "Out of Scope", "Assumptions", "Dependencies", "Change Management", "Warranty", 
        "Early Exit / Termination Clause", "Signatory (based on location)"
    ],
    "Managed Services / Support": [
        "Legal Entity", "Description of Project", "Scope", "Services Included", "Deliverables", 
        "Start and End Date", "SLA", "Payment / Invoice Schedule", "Upfront 35% Invoicing", 
        "IP or Licenses Secured/Procured by G10X", "Passthrough Expenses", "Governance Model", 
        "Out of Scope", "Assumptions", "Dependencies", "Change Management", "Warranty", 
        "Early Exit / Termination Clause", "Signatory (based on location)", "Penalty / Service Credit", 
        "Metrics, KPIs", "Capacity Fee"
    ]
}

# Collect inputs for each heading in the selected SOW type
st.subheader("Provide Information for Each Section")
section_data = {}
for heading in sow_headings[sow_type]:
    section_data[heading] = st.text_area(heading)

# Generate SOW using GPT-4
if st.button("Generate SOW"):
    if all([project_title, client_name, project_description, start_date, end_date, budget, sow_type]):
        # Construct the prompt
        prompt = (
            f"You are an expert in generating Statements of Work (SOW). Based on the following information, "
            f"create a comprehensive and professional SOW for a {sow_type} engagement:\n\n"
            f"Project Title: {project_title}\n"
            f"Client Name: {client_name}\n"
            f"Project Description: {project_description}\n"
            f"Start Date: {start_date}\n"
            f"End Date: {end_date}\n"
            f"Budget: {budget}\n\n"
        )
        
        # Add additional information based on the type of SOW selected
        if sow_type in ["T&M (Time & Materials)", "Managed Capacity"]:
            additional_info = (
                "The following roles are involved:\n"
                + "\n".join([f"- {role}: {rate}" for role, rate in roles_involved]) + "\n"
            )
            prompt += additional_info
        
        # Add section data to the prompt
        for heading in sow_headings[sow_type]:
            if heading in section_data and section_data[heading]:
                prompt += f"{heading}: {section_data[heading]}\n"

        # Generate the SOW content using OpenAI API
        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in generating Statements of Work (SOW)."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500
        )
        sow_text = response.choices[0].message['content'].strip()

        # Display the generated SOW
        st.subheader("Generated SOW")
        st.text_area("Generated SOW", sow_text, height=500)

        # Save the SOW to a Word document with tables for Timeline and Budget
        doc = Document()
        doc.add_heading(project_title, level=1)
        doc.add_heading('Client Name:', level=2)
        doc.add_paragraph(client_name)
        doc.add_heading('Project Description:', level=2)
        doc.add_paragraph(project_description)

        # Adding a table for Timeline
        doc.add_heading('Timeline:', level=2)
        timeline_table = doc.add_table(rows=2, cols=2)
        timeline_table.cell(0, 0).text = 'Start Date'
        timeline_table.cell(0, 1).text = 'End Date'
        timeline_table.cell(1, 0).text = str(start_date)
        timeline_table.cell(1, 1).text = str(end_date)

        # Adding a table for Budget
        doc.add_heading('Budget:', level=2)
        budget_table = doc.add_table(rows=1, cols=2)
        budget_table.cell(0, 0).text = 'Total Budget'
        budget_table.cell(0, 1).text = budget

        # Add additional sections based on SOW type
        if sow_type == "T&M (Time & Materials)":
            doc.add_heading('Roles & Rates:', level=2)
            for role, rate in roles_involved:
                doc.add_paragraph(f"{role}: {rate}")
        elif sow_type == "Managed Capacity":
            doc.add_heading('Unit of Work & Pricing:', level=2)
            for role, rate in roles_involved:
                doc.add_paragraph(f"{role}: {rate}")
        elif sow_type == "Fixed Bid":
            doc.add_heading('Milestones & Payment Schedule:', level=2)
            doc.add_paragraph("Milestone-based payment schedule tied to the completion of key deliverables.")
        elif sow_type == "Managed Services / Support":
            doc.add_heading('Scope & SLAs:', level=2)
            doc.add_paragraph("Scope of services, number of tickets, and SLA details with associated penalties.")

        # Add the generated SOW text
        doc.add_heading('Statement of Work:', level=2)
        doc.add_paragraph(sow_text)

        # Save to a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create download link
        st.download_button(
            label="Download SOW Document",
            data=buffer,
            file_name=f"SOW_{project_title.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please fill in all fields before generating the SOW.")